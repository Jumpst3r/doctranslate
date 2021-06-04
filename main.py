from googletrans import Translator
import argparse, googletrans
import docx
from time import sleep
import os
import string

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument('--output_language', type=str, default='fr')
    parser.add_argument('--inputfilepath', type=str)
    parser.add_argument('--outputfolder', type=str)
    args = parser.parse_args()
    if args.output_language not in googletrans.LANGCODES and args.output_language not in googletrans.LANGUAGES:
        print("invalid destination language code, must be part of")
        print(googletrans.LANGUAGES)
        exit(1)
    lang = args.output_language
    input_f = args.inputfilepath
    ouput_f = args.outputfolder + 'output.docx'
    name = '.'.join(input_f.split('.')[0:-1])
    extension = input_f.split('.')[-1]
    translator = Translator()
    if extension != 'docx':
        os.system(f'libreoffice --convert-to docx {input_f}')
    doc = docx.Document(name + '.docx')

    for i in range(len(doc.paragraphs)):
        p = doc.paragraphs[i].text
        printable = set(string.printable)
        p = ''.join(filter(lambda x: x in printable, p))
        print(p)
        res = translator.translate(p, dest=args.output_language).text
        sleep(0.5)
        doc.paragraphs[i].text = res

    doc.save(ouput_f)
    if extension != 'docx':
        os.system(f'libreoffice --headless --convert-to {extension} {ouput_f} --outdir {args.outputfolder}')
        os.system(f'rm {ouput_f}')
        
